"""Renders a one-page CV docx from a resolved CVContent structure.

This module is generation-strategy agnostic: it only knows how to lay out
already-chosen content (a single summary string, an ordered list of experience
entries with an already-chosen bullet subset, etc). Content *selection*
(tailoring) lives in generator.py; structural *variation* (section order,
phrasing choice) lives in variation.py. Both build a CVContent and hand it to
render_cv() here.
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches, RGBColor

from models import MasterCV

HEADER_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
BODY_FONT = "Calibri"
NAME_SIZE = Pt(18)
CONTACT_SIZE = Pt(9)
SECTION_HEADER_SIZE = Pt(11)
SUBHEADER_SIZE = Pt(10)
BODY_SIZE = Pt(9)

SECTION_LABELS = {
    "es": {
        "summary": "PERFIL PROFESIONAL",
        "experience": "EXPERIENCIA PROFESIONAL",
        "projects": "PROYECTO PERSONAL",
        "education": "EDUCACIÓN",
        "certifications": "CERTIFICACIONES",
        "languages": "IDIOMAS",
        "skills": "HABILIDADES E INTERESES",
        "linkedin": "Perfil de LinkedIn",
    },
    "en": {
        "summary": "PROFESSIONAL SUMMARY",
        "experience": "PROFESSIONAL EXPERIENCE",
        "projects": "PERSONAL PROJECT",
        "education": "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "languages": "LANGUAGES",
        "skills": "SKILLS & INTERESTS",
        "linkedin": "LinkedIn Profile",
    },
}


@dataclass
class ExperienceBlock:
    company: str
    role: str
    location: str
    dates: str
    bullets: list[str] = field(default_factory=list)


@dataclass
class ProjectBlock:
    name: str
    bullets: list[str] = field(default_factory=list)


@dataclass
class EducationBlock:
    institution: str
    degree_lines: list[str]  # one line per degree
    location: str
    dates: str
    notes: str = ""


@dataclass
class CVContent:
    language: str  # "es" | "en"
    name: str
    location: str
    email: str
    phone: str
    linkedin_url: str
    summary: str
    experience: list[ExperienceBlock]
    projects: list[ProjectBlock]
    education: list[EducationBlock]
    certifications: list[str]  # pre-formatted lines
    languages: list[str]  # pre-formatted lines
    skills: list[str]  # pre-formatted lines
    interests: list[str]
    interests_label: str | None = None  # e.g. "Intereses" / "Interests" — variation.py may pick an alternate label
    section_order: list[str] = field(
        default_factory=lambda: [
            "summary", "experience", "projects", "education",
            "certifications", "languages", "skills",
        ]
    )


def full_content_from_master(cv: MasterCV, language: str, summary_id: str | None = None) -> CVContent:
    """Builds a CVContent that includes *everything* in the master CV, in one
    language, with no tailoring — used to sanity-check the renderer/layout
    against the master data before any generator.py selection logic exists."""
    lang = language
    summary_variant = None
    if summary_id:
        summary_variant = next((s for s in cv.summary_bank if s.id == summary_id), None)
    if summary_variant is None:
        summary_variant = next(s for s in cv.summary_bank if s.language == lang)

    def _localize_end(end: str) -> str:
        if end.strip().lower() in ("presente", "present"):
            return "Presente" if lang == "es" else "Present"
        return end

    experience_blocks = []
    for e in cv.experience:
        role = e.role_es if lang == "es" else e.role_en
        dates = f"{e.start} – {_localize_end(e.end)}"
        bullets = [b.text_es if lang == "es" else b.text_en for b in e.bullets]
        experience_blocks.append(ExperienceBlock(e.company, role, e.location, dates, bullets))

    project_blocks = [
        ProjectBlock(p.name, [b.text_es if lang == "es" else b.text_en for b in p.bullets])
        for p in cv.projects
    ]

    education_blocks = []
    for ed in cv.education:
        degree = ed.degree_es if lang == "es" else ed.degree_en
        dates = f"{ed.start} – {ed.end}"
        notes = ed.notes_es if lang == "es" else ed.notes_en
        education_blocks.append(EducationBlock(ed.institution, [degree], ed.location, dates, notes))

    cert_lines = []
    for c in cv.certifications:
        if c.status == "in_progress":
            status = "en curso" if lang == "es" else "in progress"
            cert_lines.append(f"{c.name} — {c.issuer} — {status}, {c.date}")
        else:
            cert_lines.append(f"{c.name} — {c.issuer} — {c.date}")

    lang_lines = []
    for l in cv.languages:
        name = l.name_es if lang == "es" else l.name_en
        level = l.level_es if lang == "es" else l.level_en
        line = f"{name} — {level}"
        if l.detail:
            detail = l.detail.split(" / ")[0 if lang == "es" else -1]
            line += f" ({detail})"
        lang_lines.append(line)

    skill_lines = []
    for sg in cv.skills:
        items = sg.items_es if lang == "es" else sg.items_en
        skill_lines.append(", ".join(items))

    interests = cv.interests_es if lang == "es" else cv.interests_en

    return CVContent(
        language=lang,
        name=cv.contact.name,
        location=cv.contact.location,
        email=cv.contact.email,
        phone=cv.contact.phone,
        linkedin_url=cv.contact.linkedin_url,
        summary=summary_variant.text,
        experience=experience_blocks,
        projects=project_blocks,
        education=education_blocks,
        certifications=cert_lines,
        languages=lang_lines,
        skills=skill_lines,
        interests=list(interests),
    )


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def _add_section_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = SECTION_HEADER_SIZE
    run.font.color.rgb = HEADER_COLOR
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    borders.append(bottom)
    pPr.append(borders)


def _add_two_col_line(doc: Document, left: str, right: str, bold_left: bool = True, italic: bool = False, size: Pt = SUBHEADER_SIZE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    usable_width = Inches(8.5 - 0.6 - 0.6)
    p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(left)
    r1.bold = bold_left
    r1.italic = italic
    r1.font.size = size
    r2 = p.add_run(f"\t{right}")
    r2.italic = True
    r2.font.size = size
    return p


def _add_bullets(doc: Document, bullets: list[str]) -> None:
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        run = p.add_run(b)
        run.font.size = BODY_SIZE


def render_cv(content: CVContent) -> Document:
    doc = Document()
    _set_base_style(doc)
    labels = SECTION_LABELS[content.language]

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    name_run = name_p.add_run(content.name.upper())
    name_run.bold = True
    name_run.font.size = NAME_SIZE
    name_run.font.color.rgb = HEADER_COLOR

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(0)
    contact_run = contact_p.add_run(f"{content.location} • {content.email} • {content.phone}")
    contact_run.font.size = CONTACT_SIZE

    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_p.paragraph_format.space_after = Pt(2)
    link_run = link_p.add_run(labels["linkedin"])
    link_run.font.size = CONTACT_SIZE
    link_run.italic = True

    for section in content.section_order:
        if section == "summary" and content.summary:
            _add_section_header(doc, labels["summary"])
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(content.summary)
            r.font.size = BODY_SIZE

        elif section == "experience" and content.experience:
            _add_section_header(doc, labels["experience"])
            for e in content.experience:
                _add_two_col_line(doc, e.company, e.location, bold_left=True, size=SUBHEADER_SIZE)
                _add_two_col_line(doc, e.role, e.dates, bold_left=False, italic=True, size=BODY_SIZE)
                _add_bullets(doc, e.bullets)

        elif section == "projects" and content.projects:
            _add_section_header(doc, labels["projects"])
            for proj in content.projects:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(proj.name)
                r.bold = True
                r.font.size = SUBHEADER_SIZE
                _add_bullets(doc, proj.bullets)

        elif section == "education" and content.education:
            _add_section_header(doc, labels["education"])
            for ed in content.education:
                _add_two_col_line(doc, ed.institution, ed.location, bold_left=True, size=SUBHEADER_SIZE)
                for degree_line in ed.degree_lines:
                    _add_two_col_line(doc, degree_line, ed.dates, bold_left=False, italic=True, size=BODY_SIZE)
                if ed.notes:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    r = p.add_run(ed.notes)
                    r.italic = True
                    r.font.size = BODY_SIZE

        elif section == "certifications" and content.certifications:
            _add_section_header(doc, labels["certifications"])
            _add_bullets(doc, content.certifications)

        elif section == "languages" and content.languages:
            _add_section_header(doc, labels["languages"])
            _add_bullets(doc, content.languages)

        elif section == "skills" and (content.skills or content.interests):
            _add_section_header(doc, labels["skills"])
            lines = list(content.skills)
            if content.interests:
                interests_label = content.interests_label or ("Intereses" if content.language == "es" else "Interests")
                lines.append(f"{interests_label}: {', '.join(content.interests)}")
            _add_bullets(doc, lines)

    return doc


@dataclass
class CoverLetterContent:
    language: str
    date_line: str
    salutation: str
    paragraphs: list[str]
    closing: str
    signature_name: str


def render_cover_letter(content: CoverLetterContent) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(content.date_line)

    doc.add_paragraph()  # spacer

    sal_p = doc.add_paragraph()
    sal_p.add_run(content.salutation)

    for para_text in content.paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.add_run(para_text)

    close_p = doc.add_paragraph()
    close_p.add_run(content.closing)

    doc.add_paragraph()  # spacer before signature

    sig_p = doc.add_paragraph()
    sig_p.add_run(content.signature_name)

    return doc


def save_cover_letter(content: CoverLetterContent, output_path: str | Path) -> Path:
    doc = render_cover_letter(content)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def save_cv(content: CVContent, output_path: str | Path) -> Path:
    doc = render_cv(content)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def cv_to_bytes(content: CVContent) -> bytes:
    buf = io.BytesIO()
    render_cv(content).save(buf)
    return buf.getvalue()


def cover_letter_to_bytes(content: CoverLetterContent) -> bytes:
    buf = io.BytesIO()
    render_cover_letter(content).save(buf)
    return buf.getvalue()
